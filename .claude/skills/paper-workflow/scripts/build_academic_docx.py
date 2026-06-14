# -*- coding: utf-8 -*-
"""
将学术论文 Markdown 转换为投稿格式 DOCX。

支持学术排版：标题/副标题、作者署名、中英文摘要与关键词、作者简介、
目录、分级标题、正文、图表与题注、参考文献（含引用编号重排与交叉链接）、附录。
页面布局遵循通用中文学术期刊投稿规范（A4、宋体/Times New Roman、黑体标题）。
"""
from __future__ import annotations

import argparse
import os
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from _academic_db import normalize_reference, normalize_english_authors
from validate_citations import (
    validate as validate_citations,
    _CITE_PATTERN,
    _parse_cite_group,
)


BODY_FONT_CN = "宋体"
BODY_FONT_EN = "Times New Roman"
HEADING_FONT = "黑体"


def set_run_font(run, cn=BODY_FONT_CN, en=BODY_FONT_EN, size=12, bold=False, italic=False):
    run.font.name = en
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), cn)
    r_fonts.set(qn("w:ascii"), en)
    r_fonts.set(qn("w:hAnsi"), en)


def set_paragraph_font(paragraph, cn=BODY_FONT_CN, en=BODY_FONT_EN, size=12, bold=False):
    for run in paragraph.runs:
        set_run_font(run, cn, en, size, bold=bold)


def set_para_format(paragraph, *, align=None, first_indent=False, line_pt=20, before=0, after=0):
    fmt = paragraph.paragraph_format
    if align is not None:
        paragraph.alignment = align
    fmt.line_spacing = Pt(line_pt)
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.first_line_indent = Cm(0.74) if first_indent else None


def add_internal_hyperlink(paragraph, text, anchor, cn=BODY_FONT_CN, en=BODY_FONT_EN, size=12, superscript=False):
    run = paragraph.add_run(text)
    set_run_font(run, cn, en, size)
    run.font.superscript = superscript
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    hyperlink.set(qn("w:history"), "1")
    paragraph._p.remove(run._r)
    hyperlink.append(run._r)
    paragraph._p.append(hyperlink)


def add_inline_runs(paragraph, text, cn=BODY_FONT_CN, en=BODY_FONT_EN, size=12, bold=False, link_citations=True):
    pattern = r"(\*\*\*[^*]+\*\*\*|\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|\^\[[^\]]+\]|<sup>[^<]*</sup>)"
    for part in re.split(pattern, text):
        if not part:
            continue
        if part.startswith("***") and part.endswith("***"):
            run = paragraph.add_run(part[3:-3])
            set_run_font(run, cn, en, size, bold=True, italic=True)
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, cn, en, size, bold=True)
        elif part.startswith("*") and part.endswith("*") and not part.startswith("**"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, cn, en, size, italic=True)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, cn, "Courier New", size)
        elif part.startswith("^[") and part.endswith("]"):
            citation = part[2:-1]
            if link_citations and re.fullmatch(r"\[\d+\]", citation):
                add_internal_hyperlink(paragraph, citation, f"ref_{citation.strip('[]')}", cn, en, size, superscript=True)
            else:
                run = paragraph.add_run(citation)
                set_run_font(run, cn, en, size)
                run.font.superscript = True
        elif part.startswith("<sup>") and part.endswith("</sup>"):
            citation = part[5:-6]
            if link_citations and re.fullmatch(r"\[\d+\]", citation):
                add_internal_hyperlink(paragraph, citation, f"ref_{citation.strip('[]')}", cn, en, size, superscript=True)
            else:
                run = paragraph.add_run(citation)
                set_run_font(run, cn, en, size)
                run.font.superscript = True
        else:
            run = paragraph.add_run(part)
            set_run_font(run, cn, en, size, bold=bold)


def remove_children(element):
    for child in list(element):
        if child.tag != qn("w:sectPr"):
            element.remove(child)


def set_section_page_numbering(section, fmt="decimal", start=1):
    sect_pr = section._sectPr
    pg_num = sect_pr.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num)
    pg_num.set(qn("w:start"), str(start))
    pg_num.set(qn("w:fmt"), fmt)


def clear_paragraph(paragraph):
    p_pr = paragraph._p.pPr
    for child in list(paragraph._p):
        if child is not p_pr:
            paragraph._p.remove(child)


def add_page_number_footer(section, fmt="decimal", start=1):
    set_section_page_numbering(section, fmt=fmt, start=start)
    section.footer.is_linked_to_previous = False
    paragraph = section.footer.paragraphs[0] if section.footer.paragraphs else section.footer.add_paragraph()
    clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_begin = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run_begin._element.append(fld_begin)
    run_instr = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run_instr._element.append(instr)
    run_sep = paragraph.add_run()
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run_sep._element.append(fld_sep)
    run_text = paragraph.add_run("1")
    set_run_font(run_text, BODY_FONT_CN, BODY_FONT_EN, 10.5)
    run_end = paragraph.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run_end._element.append(fld_end)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)
    add_page_number_footer(section, fmt="lowerRoman", start=1)

    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT_EN
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT_CN)

    styles = {
        "Heading 1": (HEADING_FONT, HEADING_FONT, 16, True),
        "Heading 2": (HEADING_FONT, HEADING_FONT, 15, True),
        "Heading 3": (HEADING_FONT, HEADING_FONT, 14, True),
        "Heading 4": (HEADING_FONT, HEADING_FONT, 12, True),
    }
    for style_name, (cn, en, size, bold) in styles.items():
        try:
            style = doc.styles[style_name]
        except KeyError:
            style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = en
        style.font.size = Pt(size)
        style.font.bold = bold
        style._element.rPr.rFonts.set(qn("w:eastAsia"), cn)


def renumber_citations_and_references(text: str):
    refs = {
        int(number): body.strip()
        for number, body in re.findall(r"^\[(\d+)\]\s+(.+)$", text, flags=re.MULTILINE)
    }
    # 按正文首次出现的顺序（展开复合/区间引用后）建立 old→new 映射
    citation_order: list[int] = []
    for match in _CITE_PATTERN.finditer(text):
        inner = next(g for g in match.groups() if g is not None)
        nums, _bad = _parse_cite_group(inner)
        for n in nums:
            if n not in citation_order:
                citation_order.append(n)
    mapping = {old: new for new, old in enumerate(citation_order, start=1)}

    def replace_citation(match):
        inner = next(g for g in match.groups() if g is not None)
        nums, _bad = _parse_cite_group(inner)
        # 复合/区间一律展平为英文逗号；单引仍输出单数字
        new_inner = ",".join(str(mapping.get(n, n)) for n in nums)
        if match.group(0).startswith("<sup>"):
            return f"<sup>[{new_inner}]</sup>"
        return f"^[{new_inner}]"

    lines = text.splitlines()
    out = []
    in_refs = False
    refs_written = False
    for line in lines:
        if re.match(r"^#\s+参考文献\s*$", line.strip()):
            in_refs = True
            out.append(line)
            continue
        if in_refs:
            if re.match(r"^#\s+附录", line.strip()):
                for old in citation_order:
                    if old in refs:
                        out.append("")
                        out.append(f"[{mapping[old]}] {refs[old]}")
                out.append("")
                out.append(line)
                refs_written = True
                in_refs = False
            continue
        out.append(_CITE_PATTERN.sub(replace_citation, line))
    if in_refs and not refs_written:
        for old in citation_order:
            if old in refs:
                out.append("")
                out.append(f"[{mapping[old]}] {refs[old]}")
    report = {
        "original_citation_order": citation_order,
        "renumber_mapping": mapping,
        "removed_uncited_refs": sorted(set(refs) - set(citation_order)),
        "missing_refs": sorted(set(citation_order) - set(refs)),
    }
    return "\n".join(out), report


def parse_markdown(md_path: Path):
    text = md_path.read_text(encoding="utf-8")
    text, reference_report = renumber_citations_and_references(text)
    if text.startswith("---"):
        end = text.find("---", 3)
        if end != -1:
            text = text[end + 3 :]

    elements = []
    lines = text.splitlines()
    i = 0
    while i < len(lines):
        raw = lines[i]
        stripped = raw.strip()
        if not stripped:
            i += 1
            continue
        if stripped in {"---", "***", "___"}:
            elements.append(("sep", 0, ""))
            i += 1
            continue
        heading = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading:
            elements.append(("h", len(heading.group(1)), heading.group(2).strip()))
            i += 1
            continue
        image = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)", stripped)
        if image:
            elements.append(("img", 0, (image.group(1).strip(), image.group(2).strip())))
            i += 1
            continue
        if stripped.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            rows = []
            for row in table_lines:
                if re.match(r"^\|[\s\-:|]+\|$", row):
                    continue
                cells = [cell.strip() for cell in row.split("|")[1:-1]]
                if cells:
                    rows.append(cells)
            if rows:
                elements.append(("table", 0, rows))
            continue
        list_match = re.match(r"^(\d+[.)]|[-*+])\s+(.+)$", stripped)
        if list_match:
            marker = list_match.group(1)
            prefix = f"{marker} " if marker[0].isdigit() else "• "
            elements.append(("li", 0, prefix + list_match.group(2)))
            i += 1
            continue

        buf = [stripped]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if not nxt:
                i += 1
                break
            if (
                nxt in {"---", "***", "___"}
                or nxt.startswith("#")
                or nxt.startswith("|")
                or nxt.startswith("![")
                or re.match(r"^(\d+[.)]|[-*+])\s+.+$", nxt)
            ):
                break
            buf.append(nxt)
            i += 1
        elements.append(("p", 0, " ".join(buf)))
    return elements, reference_report


def insert_page_break(doc):
    paragraph = doc.add_paragraph()
    paragraph.add_run().add_break(WD_BREAK.PAGE)


def add_title(doc, text):
    p = doc.add_paragraph()
    set_para_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, line_pt=24, before=0, after=6)
    add_inline_runs(p, text, HEADING_FONT, HEADING_FONT, 22, bold=True)


def add_subtitle(doc, text):
    p = doc.add_paragraph()
    set_para_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, line_pt=22, before=0, after=6)
    add_inline_runs(p, text, HEADING_FONT, HEADING_FONT, 16, bold=True)


def add_author(doc, text):
    p = doc.add_paragraph()
    set_para_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, line_pt=20, after=8)
    add_inline_runs(p, text, BODY_FONT_CN, BODY_FONT_CN, 14)


def add_abstract_heading(doc, text="摘    要"):
    p = doc.add_paragraph()
    set_para_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, line_pt=20, before=8, after=8)
    add_inline_runs(p, text, HEADING_FONT, HEADING_FONT, 15, bold=True)


def add_abstract_body(doc, text, english=False):
    p = doc.add_paragraph()
    set_para_format(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line_pt=20)
    add_inline_runs(p, text, BODY_FONT_CN, BODY_FONT_EN if english else BODY_FONT_CN, 14)


def add_keywords(doc, text, english=False):
    p = doc.add_paragraph()
    set_para_format(p, align=WD_ALIGN_PARAGRAPH.LEFT, line_pt=20, before=8)
    label = "Keywords: " if english else "关键词："
    add_inline_runs(p, label, HEADING_FONT if not english else BODY_FONT_EN, BODY_FONT_EN, 14, bold=True)
    add_inline_runs(p, text, BODY_FONT_CN, BODY_FONT_EN if english else BODY_FONT_CN, 14)


def add_author_profile(doc, text, heading=False):
    p = doc.add_paragraph()
    set_para_format(
        p,
        align=WD_ALIGN_PARAGRAPH.CENTER if heading else WD_ALIGN_PARAGRAPH.LEFT,
        line_pt=20,
        before=6 if heading else 0,
    )
    add_inline_runs(p, text, HEADING_FONT if heading else BODY_FONT_CN, HEADING_FONT if heading else BODY_FONT_CN, 14, bold=heading)


def add_toc(doc):
    insert_page_break(doc)
    p = doc.add_paragraph()
    set_para_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, line_pt=20, after=8)
    add_inline_runs(p, "目    录", HEADING_FONT, HEADING_FONT, 16, bold=True)
    p = doc.add_paragraph()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    p.add_run()._element.append(fld_begin)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    p.add_run()._element.append(instr)
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    p.add_run()._element.append(fld_sep)
    run = p.add_run("请在 Word 中右键更新目录域")
    set_run_font(run, BODY_FONT_CN, BODY_FONT_EN, 12)
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    p.add_run()._element.append(fld_end)


def start_body_section(doc):
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)
    add_page_number_footer(section, fmt="decimal", start=1)
    return section


def add_heading(doc, level, text, number=None):
    style = f"Heading {min(level, 4)}"
    p = doc.add_paragraph(style=style)
    if level == 1:
        set_para_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, line_pt=20, before=12, after=8)
    else:
        set_para_format(p, align=WD_ALIGN_PARAGRAPH.LEFT, line_pt=20, before=8, after=4)
    heading_text = f"{number} {text}" if number else text
    add_inline_runs(p, heading_text, HEADING_FONT, HEADING_FONT, {1: 16, 2: 15, 3: 14}.get(level, 12), bold=True)


def next_heading_number(counters, level):
    idx = min(max(level, 1), 4) - 1
    counters[idx] += 1
    for i in range(idx + 1, len(counters)):
        counters[i] = 0
    return ".".join(str(n) for n in counters[: idx + 1] if n)


def add_bookmark(paragraph, name, bookmark_id):
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def add_body(doc, text, refs=False, appendix=False):
    p = doc.add_paragraph()
    set_para_format(
        p,
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_indent=not refs,
        line_pt=20,
        after=0,
    )
    if refs:
        p.paragraph_format.first_line_indent = None
        p.paragraph_format.left_indent = Cm(0.74)
        p.paragraph_format.first_line_indent = Cm(-0.74)
        normalized = normalize_reference(text)
        match = re.match(r"^\[(\d+)\]", normalized)
        if match:
            add_bookmark(p, f"ref_{match.group(1)}", int(match.group(1)))
        add_inline_runs(p, normalized, BODY_FONT_CN, BODY_FONT_EN, 10.5, link_citations=False)
    else:
        add_inline_runs(p, text, BODY_FONT_CN, BODY_FONT_EN, 12)


def set_cell_border(cell, **kwargs):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge, value in kwargs.items():
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        for key, val in value.items():
            element.set(qn("w:" + key), str(val))


def add_table(doc, rows):
    if not rows:
        return
    cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.text = ""
            p = cell.paragraphs[0]
            set_para_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, line_pt=20)
            add_inline_runs(p, row[c_idx] if c_idx < len(row) else "", BODY_FONT_CN, BODY_FONT_EN, 10.5, bold=(r_idx == 0))
            set_cell_border(
                cell,
                left={"val": "nil"},
                right={"val": "nil"},
                top={"val": "nil"},
                bottom={"val": "nil"},
            )
    for cell in table.rows[0].cells:
        set_cell_border(cell, top={"val": "single", "sz": 12, "color": "000000"}, bottom={"val": "single", "sz": 8, "color": "000000"})
    for cell in table.rows[-1].cells:
        set_cell_border(cell, bottom={"val": "single", "sz": 12, "color": "000000"})
    doc.add_paragraph()


def add_image(doc, md_dir: Path, alt: str, rel_path: str):
    image_path = (md_dir / rel_path).resolve()
    p = doc.add_paragraph()
    set_para_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, line_pt=20, before=8, after=4)
    if image_path.exists():
        run = p.add_run()
        run.add_picture(str(image_path), width=Cm(14.5))
    else:
        add_inline_runs(p, f"【{alt} - 图片文件未找到：{rel_path}】", BODY_FONT_CN, BODY_FONT_EN, 10.5)


def add_caption(doc, text):
    clean = re.sub(r"^\*\*|\*\*$", "", text.strip())
    p = doc.add_paragraph()
    set_para_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, line_pt=20, after=6)
    add_inline_runs(p, clean, BODY_FONT_CN, BODY_FONT_CN, 10.5, bold=True)


def build_docx(md_path: Path, out_path: Path, template_path: Path | None = None):
    doc = Document(str(template_path)) if template_path else Document()
    if template_path:
        remove_children(doc.element.body)
    configure_document(doc)

    elements, reference_report = parse_markdown(md_path)
    md_text, _ = renumber_citations_and_references(md_path.read_text(encoding="utf-8"))
    citation_report = validate_citations(md_text)
    citation_report.update(reference_report)

    area = "front"
    title_seen = 0
    skip_manual_toc = False
    toc_added = False
    body_section_started = False
    heading_counters = [0, 0, 0, 0]
    pending_table_caption = None

    for kind, level, payload in elements:
        if skip_manual_toc:
            if kind == "h" and level == 1 and payload != "附录 1":
                skip_manual_toc = False
            else:
                continue

        if kind == "sep":
            continue

        if kind == "h":
            text = payload
            if "目录" in text:
                skip_manual_toc = True
                continue
            if "参考文献" in text:
                insert_page_break(doc)
                area = "refs"
                add_heading(doc, 1, "参考文献")
                continue
            if text.startswith("附录"):
                area = "appendix"
                add_heading(doc, 1, text)
                continue
            if area == "appendix":
                add_heading(doc, level, text)
                continue
            area = "body"
            if not toc_added:
                add_toc(doc)
                toc_added = True
            if not body_section_started:
                start_body_section(doc)
                body_section_started = True
            number = next_heading_number(heading_counters, level)
            add_heading(doc, level, text, number=number)
            continue

        if kind == "img":
            alt, rel = payload
            add_image(doc, md_path.parent, alt, rel)
            continue

        if kind == "table":
            if pending_table_caption:
                add_caption(doc, pending_table_caption)
                pending_table_caption = None
            add_table(doc, payload)
            continue

        if kind == "li":
            add_body(doc, payload, refs=False)
            continue

        text = payload.strip()
        if not text:
            continue

        if area == "front":
            if title_seen == 0:
                add_title(doc, text)
                title_seen += 1
                continue
            if text.startswith("——"):
                add_subtitle(doc, text)
                continue
            if text.startswith("作者姓名"):
                add_author(doc, text)
                continue
            if re.match(r"^【摘\s*要】", text):
                add_abstract_heading(doc)
                area = "abstract_cn"
                continue

        if area == "abstract_cn":
            if text.startswith("【关键词"):
                keywords = re.sub(r"^【关键词】[:：]?", "", text).strip()
                add_keywords(doc, keywords)
                continue
            if re.match(r"^【Abstract】", text):
                insert_page_break(doc)
                add_abstract_heading(doc, "Abstract")
                area = "abstract_en"
                continue
            add_abstract_body(doc, text)
            continue

        if area == "abstract_en":
            if text.startswith("【Keywords】") or text.startswith("【Keywords】:"):
                keywords = re.sub(r"^【Keywords】[:：]?", "", text).strip()
                add_keywords(doc, keywords, english=True)
                continue
            if text == "作者简介":
                insert_page_break(doc)
                area = "author_profile"
                add_author_profile(doc, text, heading=True)
                continue
            add_abstract_body(doc, text, english=True)
            continue

        if area == "author_profile":
            if text == "作者简介":
                add_author_profile(doc, text, heading=True)
            else:
                add_author_profile(doc, text)
            continue

        if re.match(r"^\*\*图\s*\d+", text):
            add_caption(doc, text)
            continue

        if re.match(r"^\*\*表\s", text):
            pending_table_caption = text
            continue

        if area == "refs":
            add_body(doc, text, refs=True)
            continue

        add_body(doc, text, refs=False, appendix=(area == "appendix"))

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return citation_report


def _preflight_validate(md_path: Path) -> None:
    """build 前自检：bad_tokens / missing_refs 任一非空即 abort。

    这是 SKILL.md "build 前置 validate" 的代码契约化——避免作者跳过 validate
    把含编号问题的 md 直接 build 成投稿稿件。
    """
    md_text = md_path.read_text(encoding="utf-8")
    result = validate_citations(md_text)
    fatal = []
    if result["bad_tokens"]:
        fatal.append(f"无法解析的引用 token: {sorted(set(result['bad_tokens']))}")
    if result["missing_refs"]:
        fatal.append(f"缺失参考文献定义: {result['missing_refs']}")
    if fatal:
        msg = "\n  - ".join(["build 前置自检不通过，已中止："] + fatal)
        msg += "\n（请先跑 validate_citations.py 修正后再 build）"
        raise SystemExit(msg)


def main():
    # Windows 控制台 utf-8 输出，与 validate_citations 一致
    try:
        import sys
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")
        sys.stderr.reconfigure(encoding="utf-8", errors="replace")
    except (AttributeError, ImportError):
        pass
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", required=True, type=Path)
    parser.add_argument("--output", required=True, type=Path)
    parser.add_argument("--template", type=Path)
    args = parser.parse_args()
    _preflight_validate(args.input.resolve())
    report = build_docx(args.input.resolve(), args.output.resolve(), args.template.resolve() if args.template else None)
    print(f"完成: {args.output.resolve()} ({os.path.getsize(args.output):,} bytes)")
    print("引用编号:", report)


if __name__ == "__main__":
    main()
